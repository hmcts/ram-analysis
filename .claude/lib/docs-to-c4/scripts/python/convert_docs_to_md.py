"""Convert binary document files (.docx, .pdf, .xlsx) to Markdown.

Recursively scans an input folder for supported binary formats, extracts
their text content, and writes .md files to an output folder preserving
the relative directory structure.

Supported formats:
  .docx  — paragraphs, tables, headings (via python-docx)
  .pdf   — page-by-page text extraction (via pypdf)
  .xlsx  — sheet-by-sheet table extraction (via openpyxl)

Files that are already text-based (.md, .txt, .csv, .json, .yaml, .yml)
are copied as-is to the output folder.

Usage:
  python convert_docs_to_md.py <input_folder> <output_folder>
  python convert_docs_to_md.py <input_folder> <output_folder> --skip-existing
"""

from __future__ import annotations

import argparse
import json
import shutil
import sys
from pathlib import Path

BINARY_EXTENSIONS = {".docx", ".pdf", ".xlsx"}
TEXT_EXTENSIONS = {".md", ".txt", ".csv", ".json", ".yaml", ".yml"}
SUPPORTED_EXTENSIONS = BINARY_EXTENSIONS | TEXT_EXTENSIONS

SKIP_DIRS = {
    "node_modules", ".git", "__pycache__", ".venv", "venv",
    ".claude", "_bmad-output", ".cursor", ".vscode", "output",
}

CHARS_PER_TOKEN = 4


def extract_docx(path: Path) -> str:
    """Extract text from a .docx file as Markdown."""
    from docx import Document
    from docx.table import Table

    doc = Document(str(path))
    lines: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Find the matching paragraph object
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()
                    if not text:
                        break
                    style = (para.style.name or "").lower()
                    if "heading 1" in style:
                        lines.append(f"# {text}")
                    elif "heading 2" in style:
                        lines.append(f"## {text}")
                    elif "heading 3" in style:
                        lines.append(f"### {text}")
                    elif "heading 4" in style:
                        lines.append(f"#### {text}")
                    elif "list" in style:
                        lines.append(f"- {text}")
                    else:
                        lines.append(text)
                    break

        elif tag == "tbl":
            for table in doc.tables:
                if table._element is element:
                    lines.append(_format_table(table))
                    break

    return "\n\n".join(lines)


def _format_table(table) -> str:
    """Convert a docx Table to a Markdown table."""
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    col_count = max(len(r) for r in rows)
    # Pad rows to uniform width
    for r in rows:
        while len(r) < col_count:
            r.append("")

    md_lines = []
    # Header row
    md_lines.append("| " + " | ".join(rows[0]) + " |")
    md_lines.append("| " + " | ".join(["---"] * col_count) + " |")
    # Data rows
    for row in rows[1:]:
        md_lines.append("| " + " | ".join(row) + " |")

    return "\n".join(md_lines)


def extract_pdf(path: Path) -> str:
    """Extract text from a .pdf file, one section per page."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    sections: list[str] = []

    for i, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        if text:
            sections.append(f"## Page {i}\n\n{text}")

    return "\n\n".join(sections)


def extract_xlsx(path: Path) -> str:
    """Extract data from an .xlsx file as Markdown tables per sheet."""
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    sections: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() if c is not None else "" for c in row]
            if any(cells):
                rows.append(cells)

        if not rows:
            continue

        col_count = max(len(r) for r in rows)
        for r in rows:
            while len(r) < col_count:
                r.append("")

        md_lines = [f"## {sheet_name}"]
        md_lines.append("| " + " | ".join(rows[0]) + " |")
        md_lines.append("| " + " | ".join(["---"] * col_count) + " |")
        for row in rows[1:]:
            md_lines.append("| " + " | ".join(row) + " |")

        sections.append("\n".join(md_lines))

    wb.close()
    return "\n\n".join(sections)


EXTRACTORS = {
    ".docx": extract_docx,
    ".pdf": extract_pdf,
    ".xlsx": extract_xlsx,
}


def find_files(input_folder: Path) -> list[Path]:
    """Recursively find supported files, skipping excluded directories."""
    files: list[Path] = []
    for item in sorted(input_folder.rglob("*")):
        if any(skip in item.parts for skip in SKIP_DIRS):
            continue
        if item.is_file() and item.suffix.lower() in SUPPORTED_EXTENSIONS:
            files.append(item)
    return files


def convert_file(src: Path, dest: Path) -> dict:
    """Convert a single file. Returns a result dict."""
    dest.parent.mkdir(parents=True, exist_ok=True)
    ext = src.suffix.lower()
    original_size = src.stat().st_size

    if ext in TEXT_EXTENSIONS:
        shutil.copy2(src, dest)
        converted_size = dest.stat().st_size
        return {
            "source": str(src),
            "output": str(dest),
            "action": "copied",
            "original_bytes": original_size,
            "output_bytes": converted_size,
            "estimated_tokens": converted_size // CHARS_PER_TOKEN,
        }

    extractor = EXTRACTORS.get(ext)
    if not extractor:
        return {
            "source": str(src),
            "output": None,
            "action": "skipped",
            "reason": f"unsupported format: {ext}",
        }

    try:
        md_content = extractor(src)
        md_dest = dest.with_suffix(".md")
        md_dest.write_text(md_content, encoding="utf-8")
        converted_size = md_dest.stat().st_size
        return {
            "source": str(src),
            "output": str(md_dest),
            "action": "converted",
            "original_bytes": original_size,
            "output_bytes": converted_size,
            "estimated_tokens": converted_size // CHARS_PER_TOKEN,
        }
    except Exception as e:
        return {
            "source": str(src),
            "output": None,
            "action": "error",
            "reason": str(e),
        }


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert binary document files to Markdown.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("input_folder", help="Folder containing source documents")
    parser.add_argument("output_folder", help="Folder to write converted .md files")
    parser.add_argument(
        "--skip-existing",
        action="store_true",
        help="Skip files whose output already exists",
    )
    args = parser.parse_args()

    input_folder = Path(args.input_folder).resolve()
    output_folder = Path(args.output_folder).resolve()

    if not input_folder.is_dir():
        print(json.dumps({"status": "error", "error": f"Not a directory: {input_folder}"}))
        sys.exit(1)

    output_folder.mkdir(parents=True, exist_ok=True)

    files = find_files(input_folder)
    if not files:
        print(json.dumps({"status": "error", "error": "No supported files found"}))
        sys.exit(1)

    results = []
    total_original = 0
    total_converted = 0

    for src in files:
        rel = src.relative_to(input_folder)
        dest = output_folder / rel

        if args.skip_existing:
            md_dest = dest.with_suffix(".md") if src.suffix.lower() in BINARY_EXTENSIONS else dest
            if md_dest.exists():
                results.append({
                    "source": str(src),
                    "output": str(md_dest),
                    "action": "skipped_existing",
                })
                continue

        result = convert_file(src, dest)
        results.append(result)

        if result.get("original_bytes"):
            total_original += result["original_bytes"]
        if result.get("output_bytes"):
            total_converted += result["output_bytes"]

    converted_count = sum(1 for r in results if r["action"] == "converted")
    copied_count = sum(1 for r in results if r["action"] == "copied")
    error_count = sum(1 for r in results if r["action"] == "error")

    report = {
        "status": "ok",
        "input_folder": str(input_folder),
        "output_folder": str(output_folder),
        "files": results,
        "summary": {
            "total_files": len(files),
            "converted": converted_count,
            "copied": copied_count,
            "errors": error_count,
            "original_total_bytes": total_original,
            "converted_total_bytes": total_converted,
            "estimated_total_tokens": total_converted // CHARS_PER_TOKEN,
        },
    }

    print(json.dumps(report, indent=2))


if __name__ == "__main__":
    main()
